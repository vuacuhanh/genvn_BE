from typing import Dict, Any, List
import io

print("LOADED exporter.py pretty")

# =========================
# Helpers chung
# =========================
def _get_meta(paper: Dict[str, Any]) -> Dict[str, Any]:
    return paper.get("meta", {}) or {}

def _iter_items(paper: Dict[str, Any]) -> List[Dict[str, Any]]:
    return paper.get("items", []) or []

def _find_writing_items(paper: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [it for it in _iter_items(paper) if it.get("type") == "writing"]


# =========================
# DOCX EXPORT
# =========================
def paper_to_docx(paper: Dict[str, Any]) -> bytes:
    """
    Tạo DOCX có:
      - Header: Họ tên / Lớp / Ngày
      - Tiêu đề + Theme + Thời gian làm bài
      - Khung Điểm / Nhận xét (table)
      - MCQ hiển thị 4 phương án
      - Writing: chèn các dòng trống để HS viết
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
    except ImportError as e:
        raise RuntimeError("Thiếu dependency `python-docx`. Cài: pip install python-docx") from e

    doc = Document()
    meta = _get_meta(paper)

    # -------- Header thông tin HS
    p_info = doc.add_paragraph()
    run = p_info.add_run("Họ tên: ____________________    Lớp: ________    Ngày: ____/____/______")
    run.font.size = Pt(11)

    # -------- Tiêu đề
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ĐỀ TIẾNG VIỆT - LỚP {meta.get('grade','')}")
    run.bold = True
    run.font.size = Pt(18)

    # -------- Theme + Thời gian
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    theme_txt = meta.get("theme") or "—"
    time_txt = meta.get("estimated_time_minutes") or ""
    p_meta.add_run(f"Chủ điểm: {theme_txt}").italic = True
    if time_txt:
        p_meta.add_run(f"   •   Thời gian làm bài: {time_txt} phút")

    doc.add_paragraph("")  # spacing

    # -------- Khung Điểm & Nhận xét (table 1x2 + 2x1)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(10)

    table.cell(0, 0).text = "Điểm:"
    table.cell(0, 1).text = "Nhận xét:"
    table.cell(1, 0).text = "………………"
    table.cell(1, 1).text = "………………………………………………………………………………\n………………………………………………………………………………"

    doc.add_paragraph("")  # spacing

    # -------- Nội dung câu hỏi
    items = _iter_items(paper)
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"Câu {idx} ({item.get('score',1)}đ) [{item.get('type')}]: ")
        run.bold = True
        p.add_run(item.get("prompt", ""))

        # MCQ
        if item.get("type") == "mcq":
            opts = item.get("options", [])
            for label, opt in zip(["A", "B", "C", "D"], opts):
                para = doc.add_paragraph(style=None)
                para.paragraph_format.left_indent = Cm(0.5)
                para.add_run(f"{label}. {opt}")

        # Writing: chèn dòng trống để viết
        if item.get("type") == "writing":
            # số dòng viết: nếu không có gợi ý length thì tạo 8 dòng
            lines = 8
            # tạo các dòng bằng chuỗi gạch chấm cho đơn giản & đẹp
            for _ in range(lines):
                doc.add_paragraph("........................................................................................")

        doc.add_paragraph("")  # spacing mỗi câu

    # -------- Đáp án
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run("ĐÁP ÁN")
    r.bold = True
    r.font.size = Pt(14)

    for ans in paper.get("answer_sheet", []):
        doc.add_paragraph(f"{ans.get('id')}: {ans.get('key')}")

    # (Optional) Watermark trong Word rất khó bằng python-docx;
    # nếu cần watermark ở DOCX, nên dùng PDF để đảm bảo hiển thị đồng nhất.

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()


# =========================
# TXT EXPORT (giữ nguyên, gọn)
# =========================
def paper_to_txt(paper: Dict[str, Any]) -> str:
    meta = _get_meta(paper)
    lines = []
    lines.append(f"DE TIENG VIET - LOP {meta.get('grade','')}")
    lines.append(f"Do kho: {meta.get('difficulty','')}")
    if meta.get("theme"):
        lines.append(f"Chu diem: {meta.get('theme')}")
    if meta.get("estimated_time_minutes"):
        lines.append(f"Thoi gian: {meta.get('estimated_time_minutes')} phut")
    lines.append("")
    lines.append("Ho ten: _____________   Lop: ________   Ngay: ____/____/______")
    lines.append("Diem: ________   Nhan xet: __________________________________")
    lines.append("")

    for idx, item in enumerate(_iter_items(paper), start=1):
        lines.append(f"Cau {idx} ({item.get('score',1)}d) [{item.get('type')}]: {item.get('prompt','')}")
        if item.get("type") == "mcq":
            for opt_label, opt in zip(["A","B","C","D"], item.get("options", [])):
                lines.append(f"  {opt_label}. {opt}")
        if item.get("type") == "writing":
            lines += ["  ............................................................." for _ in range(8)]
        lines.append("")

    lines.append("Dap an:")
    for ans in paper.get("answer_sheet", []):
        lines.append(f"- {ans.get('id')}: {ans.get('key')}")
    return "\n".join(lines)


# =========================
# PDF EXPORT (có watermark)
# =========================
def paper_to_pdf(paper: Dict[str, Any]) -> bytes:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    FONT_PATH = os.path.join(os.path.dirname(__file__), "../fonts/InterTight-Italic-VariableFont_wght.ttf")
    pdfmetrics.registerFont(TTFont("InterTight", FONT_PATH))

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas as pdfcanvas

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
        title="Đề Tiếng Việt"
    )

    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "InterTight"
    styles["Title"].fontName = "InterTight"

    styles.add(ParagraphStyle(name="Q", fontName="InterTight", fontSize=11, leading=15))
    styles.add(ParagraphStyle(name="Small", fontName="InterTight", fontSize=9, leading=12))

    story = []
    meta = _get_meta(paper)

    # Header: Họ tên / Lớp / Ngày
    story.append(Paragraph("Họ tên: ____________________    Lớp: ________    Ngày: ____/____/______", styles["Small"]))
    story.append(Spacer(1, 6))

    # Title
    story.append(Paragraph(f"<b>ĐỀ TIẾNG VIỆT - LỚP {meta.get('grade','')}</b>", styles["Title"]))

    # Theme + Time
    theme_txt = meta.get("theme") or "—"
    extra = []
    extra.append(f"Chủ điểm: {theme_txt}")
    if meta.get("estimated_time_minutes"):
        extra.append(f"Thời gian: {meta.get('estimated_time_minutes')} phút")
    story.append(Paragraph("   •   ".join(extra), styles["Small"]))
    story.append(Spacer(1, 8))

    # Khung điểm / nhận xét: Table (2 hàng 2 cột)
    box = Table(
        data=[
            ["Điểm:", "Nhận xét:"],
            ["………………", "………………………………………………………………………………\n………………………………………………………………………………"]
        ],
        colWidths=[5 * cm, 10 * cm]
    )
    box.setStyle(TableStyle([
        ("BOX", (0,0), (-1,-1), 0.5, colors.grey),
        ("INNERGRID", (0,0), (-1,-1), 0.25, colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
        ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ("TOPPADDING", (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]))
    story.append(box)
    story.append(Spacer(1, 10))

    # Câu hỏi
    for idx, item in enumerate(_iter_items(paper), start=1):
        story.append(Paragraph(f"<b>Câu {idx}</b> ({item.get('score',1)}đ) [{item.get('type')}]: {item.get('prompt','')}", styles["Q"]))

        if item.get("type") == "mcq":
            opts = item.get("options", [])
            for label, opt in zip(["A", "B", "C", "D"], opts):
                story.append(Paragraph(f"{label}. {opt}", styles["Q"]))

        if item.get("type") == "writing":
            # kẻ 8 dòng để viết
            story.append(Spacer(1, 4))
            lines = []
            for _ in range(8):
                # vẽ dòng bằng một table mỏng (đẹp & thẳng)
                t = Table([[" "]], colWidths=[16 * cm], rowHeights=[0.6 * cm])
                t.setStyle(TableStyle([
                    ("LINEBELOW", (0,0), (-1,-1), 0.5, colors.black),
                ]))
                lines.append(t)
            story.extend(lines)

        story.append(Spacer(1, 6))

    # Trang đáp án
    story.append(PageBreak())
    story.append(Paragraph("<b>ĐÁP ÁN</b>", styles["Heading2"]))
    for ans in paper.get("answer_sheet", []):
        story.append(Paragraph(f"{ans.get('id')}: {ans.get('key')}", styles["Q"]))

    # ---- Watermark callback
    def _add_watermark(canvas: pdfcanvas.Canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 48)
        canvas.setFillColorRGB(0.85, 0.85, 0.85)
        canvas.translate(300, 400)
        canvas.rotate(30)
        canvas.drawCentredString(0, 0, "VN PRIMARY GENERATOR")
        canvas.restoreState()

    doc.build(story, onFirstPage=_add_watermark, onLaterPages=_add_watermark)
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data
